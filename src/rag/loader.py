import json
import logging
import re
import unicodedata
from pathlib import Path
from typing import Iterable, List, Tuple

from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from tqdm import tqdm

from src.rag.fee_fact_chunker import extract_fee_fact_documents

def clean_vietnamese_text(text: str, preserve_newlines: bool = True) -> str:
    """Normalize and clean Vietnamese text extracted from PDFs."""
    text = unicodedata.normalize("NFC", text)

    text = "".join(
        char
        for char in text
        if not unicodedata.category(char).startswith("C") or char in "\n\t"
    )

    text = text.replace("\r\n", "\n").replace("\r", "\n")
    if preserve_newlines:
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n[ \t]+", "\n", text)
        text = re.sub(r"\n{2,}", "\n", text)
    else:
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"\n\s*\n", "\n", text)

    return text.strip()


logger = logging.getLogger(__name__)


class SimpleLoader:
    """Load documents from a directory with basic format support."""

    def __init__(
        self,
        include_tables: bool = True,
        table_rows_per_chunk: int = 8,
        fee_fact_chunking: bool = True,
        fee_fact_patterns: Iterable[str] | None = None,
        fee_fact_include_raw_text: bool = True,
    ):
        self.include_tables = include_tables
        self.table_rows_per_chunk = table_rows_per_chunk
        self.fee_fact_chunking = fee_fact_chunking
        self.fee_fact_include_raw_text = fee_fact_include_raw_text
        self.fee_fact_patterns = (
            tuple(p.lower() for p in fee_fact_patterns)
            if fee_fact_patterns
            else ("bieu_phi", "bieu-phi", "fee")
        )

    def _is_fee_fact_pdf(self, pdf_file: Path) -> bool:
        name = pdf_file.name.lower()
        return any(pattern in name for pattern in self.fee_fact_patterns)

    def load_pdf(self, pdf_file: Path) -> List[Document]:
        fee_fact_docs: List[Document] = []
        if self.fee_fact_chunking and self._is_fee_fact_pdf(pdf_file):
            try:
                fee_fact_docs = extract_fee_fact_documents(pdf_file)
            except ImportError:
                logger.warning("pdfplumber not installed; skipping fee-fact chunking.")
            except Exception as exc:
                logger.warning("fee-fact chunking failed for %s: %s", pdf_file, exc)

        docs: List[Document] = []
        if self.fee_fact_include_raw_text or not fee_fact_docs:
            docs = PyPDFLoader(str(pdf_file), extract_images=True).load()
            for doc in docs:
                doc.page_content = clean_vietnamese_text(
                    doc.page_content, preserve_newlines=True
                )
                doc.metadata.setdefault("source", str(pdf_file))
                doc.metadata.setdefault("file_type", "pdf")

        if self.include_tables and not fee_fact_docs:
            table_docs = self._load_pdf_tables(pdf_file)
            if not table_docs:
                table_docs = self._load_text_tables(docs, pdf_file)
            docs.extend(table_docs)

        if fee_fact_docs:
            docs.extend(fee_fact_docs)

        return docs

    def _load_pdf_tables(self, pdf_file: Path) -> List[Document]:
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber not installed; skipping table extraction.")
            return []

        table_docs: List[Document] = []
        with pdfplumber.open(str(pdf_file)) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                tables = page.extract_tables() or []
                for table_index, table in enumerate(tables, start=1):
                    table_docs.extend(
                        self._table_to_json_docs(
                            table,
                            metadata_base={
                                "source": str(pdf_file),
                                "file_type": "pdf_table",
                                "page": page_index,
                                "table_index": table_index,
                            },
                        )
                    )
        return table_docs

    def _load_text_tables(
        self, docs: List[Document], pdf_file: Path
    ) -> List[Document]:
        table_docs: List[Document] = []
        for doc in docs:
            lines = [line.strip() for line in doc.page_content.splitlines() if line.strip()]
            header_lines, rows = self._extract_table_rows(lines)
            if not rows:
                continue

            table_docs.extend(
                self._table_rows_to_json_docs(
                    rows,
                    header_lines=header_lines,
                    metadata_base={
                        "source": str(pdf_file),
                        "file_type": "pdf_table_text",
                        "page": doc.metadata.get("page", None),
                        "is_table": True,
                    },
                )
            )
        return table_docs

    def _extract_table_rows(
        self, lines: List[str]
    ) -> Tuple[List[str], List[dict]]:
        header_lines: List[str] = []
        rows: List[dict] = []
        current_row: List[str] = []
        current_row_id: str | None = None
        current_section: str | None = None

        header_keywords = [
            "stt",
            "ma",
            "mã",
            "phi",
            "phí",
            "dich vu",
            "dịch vụ",
            "muc phi",
            "mức phí",
            "toi thieu",
            "tối thiểu",
            "toi da",
            "tối đa",
        ]

        def is_header(line: str) -> bool:
            lowered = line.lower()
            score = sum(1 for keyword in header_keywords if keyword in lowered)
            if score < 2:
                return False
            return not is_row_start(line)

        def is_section(line: str) -> bool:
            return bool(re.match(r"^[IVX]+\b", line)) and not re.search(r"\d", line)

        def is_row_start(line: str) -> bool:
            return bool(
                re.match(r"^\d+(?:\.\d+)?\b", line)
                or re.match(r"^[A-Z]{1,3}-?\d", line)
            )

        def flush_row():
            nonlocal current_row, current_row_id
            if not current_row:
                return
            text = " ".join(current_row).strip()
            if text:
                rows.append(
                    {
                        "row_id": current_row_id or "",
                        "text": text,
                        "section": current_section,
                    }
                )
            current_row = []
            current_row_id = None

        for line in lines:
            if is_section(line):
                flush_row()
                current_section = line
                continue

            if is_header(line) and not rows and not current_row:
                header_lines.append(line)
                continue

            if is_row_start(line):
                flush_row()
                match = re.match(r"^(\d+(?:\.\d+)?\b)", line)
                if match:
                    current_row_id = match.group(1)
                    rest = line[match.end() :].strip()
                    current_row = [rest] if rest else []
                else:
                    current_row_id = ""
                    current_row = [line]
                continue

            if current_row:
                current_row.append(line)

        flush_row()

        if len(rows) < 2:
            return header_lines, []
        return header_lines, rows

    def _table_rows_to_json_docs(
        self,
        rows: List[dict],
        header_lines: List[str],
        metadata_base: dict,
    ) -> List[Document]:
        docs: List[Document] = []
        for chunk_index, row_group in enumerate(self._chunk_rows(rows), start=1):
            payload = {
                "type": "table_text",
                "columns": ["row_id", "text", "section"],
                "rows": row_group,
                "header_lines": header_lines,
                "row_count": len(row_group),
                "header_detected": bool(header_lines),
                "source": "heuristic",
            }
            metadata = dict(metadata_base)
            metadata.update(
                {
                    "is_table": True,
                    "chunk_index": chunk_index,
                }
            )
            docs.append(
                Document(
                    page_content=json.dumps(payload, ensure_ascii=False),
                    metadata=metadata,
                )
            )
        return docs

    def _normalize_rows(self, table: List[List[str | None]]) -> List[List[str]]:
        if not table:
            return []

        max_cols = max((len(row) for row in table if row), default=0)
        normalized = []
        for row in table:
            row = row or []
            padded = [(cell or "").strip() for cell in row]
            padded.extend([""] * (max_cols - len(padded)))
            if any(cell.strip() for cell in padded):
                normalized.append(padded)
        return normalized

    def _is_header_row(self, row: List[str]) -> bool:
        row_text = " ".join(row).lower()
        header_keywords = [
            "stt",
            "ma",
            "phi",
            "dich vu",
            "muc phi",
            "toi thieu",
            "toi da",
            "kh",
        ]
        if any(keyword in row_text for keyword in header_keywords):
            return True

        alpha_cells = sum(1 for cell in row if sum(ch.isalpha() for ch in cell) >= 2)
        numeric_cells = sum(
            1 for cell in row if cell.replace(",", "").replace(".", "").isdigit()
        )
        return alpha_cells >= max(2, numeric_cells)

    def _split_header(
        self, rows: List[List[str]]
    ) -> Tuple[List[str], List[List[str]], bool]:
        if not rows:
            return [], [], False

        header = rows[0]
        has_header = self._is_header_row(header)
        if has_header:
            columns = header
            data_rows = rows[1:]
        else:
            columns = [f"col_{idx + 1}" for idx in range(len(header))]
            data_rows = rows

        cleaned_columns = []
        seen = set()
        for idx, col in enumerate(columns):
            name = (col or "").strip() or f"col_{idx + 1}"
            if name in seen:
                suffix = 2
                while f"{name}_{suffix}" in seen:
                    suffix += 1
                name = f"{name}_{suffix}"
            cleaned_columns.append(name)
            seen.add(name)

        return cleaned_columns, data_rows, has_header

    def _chunk_rows(self, rows: List[List[str]]) -> List[List[List[str]]]:
        if self.table_rows_per_chunk <= 0:
            return [rows]
        return [
            rows[i : i + self.table_rows_per_chunk]
            for i in range(0, len(rows), self.table_rows_per_chunk)
        ]

    def _table_to_json_docs(
        self, table: List[List[str | None]], metadata_base: dict
    ) -> List[Document]:
        normalized_rows = self._normalize_rows(table)
        if not normalized_rows:
            return []

        columns, data_rows, has_header = self._split_header(normalized_rows)
        if not columns:
            return []

        docs: List[Document] = []
        for chunk_index, row_group in enumerate(self._chunk_rows(data_rows), start=1):
            rows_payload = [dict(zip(columns, row)) for row in row_group]
            payload = {
                "type": "table",
                "columns": columns,
                "rows": rows_payload,
                "header_detected": has_header,
                "row_count": len(rows_payload),
            }
            metadata = dict(metadata_base)
            metadata.update(
                {
                    "is_table": True,
                    "chunk_index": chunk_index,
                }
            )
            docs.append(
                Document(
                    page_content=json.dumps(payload, ensure_ascii=False),
                    metadata=metadata,
                )
            )
        return docs

    def load_text(self, text_file: Path, file_type: str) -> List[Document]:
        content = text_file.read_text(encoding="utf-8")
        content = clean_vietnamese_text(content)
        return [
            Document(
                page_content=content,
                metadata={"source": str(text_file), "file_type": file_type},
            )
        ]

    def load_docx(self, docx_file: Path) -> List[Document]:
        try:
            import docx
        except ImportError as exc:
            raise ImportError("python-docx is required for .docx files") from exc

        doc = docx.Document(str(docx_file))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        content = clean_vietnamese_text("\n".join(paragraphs))
        return [
            Document(
                page_content=content,
                metadata={"source": str(docx_file), "file_type": "docx"},
            )
        ]

    def load_xlsx(self, xlsx_file: Path) -> List[Document]:
        try:
            import pandas as pd
        except ImportError as exc:
            raise ImportError("pandas is required for .xlsx files") from exc

        sheets = pd.read_excel(xlsx_file, sheet_name=None)
        docs: List[Document] = []
        for sheet_name, df in sheets.items():
            table = [df.columns.tolist()] + df.fillna("").astype(str).values.tolist()
            docs.extend(
                self._table_to_json_docs(
                    table,
                    metadata_base={
                        "source": str(xlsx_file),
                        "file_type": "xlsx_table",
                        "sheet": sheet_name,
                        "is_table": True,
                    },
                )
            )
        return docs

    def load_dir(self, dir_path: Path, file_types: Iterable[str] | None = None) -> List:
        dir_path = Path(dir_path)
        if file_types is None:
            file_types = ("pdf",)
        file_types = tuple(ext.lower().lstrip(".") for ext in file_types)

        files = []
        for ext in file_types:
            files.extend(sorted(dir_path.glob(f"*.{ext}")))

        if not files:
            raise ValueError(f"No files found in {dir_path} for {file_types}")

        all_docs: List[Document] = []
        for file_path in tqdm(files, desc="Loading documents"):
            try:
                ext = file_path.suffix.lower().lstrip(".")
                if ext == "pdf":
                    all_docs.extend(self.load_pdf(file_path))
                elif ext in {"txt", "md"}:
                    all_docs.extend(self.load_text(file_path, ext))
                elif ext == "docx":
                    all_docs.extend(self.load_docx(file_path))
                elif ext == "xlsx":
                    all_docs.extend(self.load_xlsx(file_path))
            except Exception:
                continue

        return all_docs
